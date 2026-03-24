"""
Generate Word-style deliverables (.docx) for the volatility prediction project.
Run from project root: python scripts/build_word_docs.py
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

DOCS_DIR = ROOT / "docs"
METRICS_PATH = ROOT / "artifacts" / "metrics.json"


def _load_metrics() -> dict:
    if METRICS_PATH.exists():
        with open(METRICS_PATH, encoding="utf-8") as f:
            return json.load(f)
    return {}


def _set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _title(doc: Document, text: str) -> None:
    h = doc.add_heading(text, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def _h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def _p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def _bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def build_hld() -> Document:
    doc = Document()
    _set_normal_style(doc)
    _title(doc, "High-Level Design (HLD)")
    _p(doc, "Cryptocurrency Volatility Prediction System")
    _p(doc, "Document version: 1.0")

    _h1(doc, "1. Purpose and scope")
    _p(
        doc,
        "This system predicts short-horizon realized volatility for cryptocurrencies "
        "using daily market data (OHLC, volume, market capitalization). The goal is to "
        "support risk awareness and decision-making by estimating how volatile the next "
        "few trading days may be, based on historical patterns.",
    )

    _h1(doc, "2. Stakeholders and users")
    _bullets(
        doc,
        [
            "Analysts and students evaluating ML workflows on crypto time series.",
            "Operators running batch training and reviewing metrics locally.",
            "End users testing predictions through a simple Streamlit web interface.",
        ],
    )

    _h1(doc, "3. System overview")
    _p(
        doc,
        "The solution is a batch machine learning pipeline plus a local interactive app. "
        "Raw data lives in a CSV file. Python modules load and clean the data, engineer "
        "features and the volatility target, train a gradient boosting regressor, persist "
        "the model and metrics, and optionally serve predictions through Streamlit.",
    )

    _h1(doc, "4. Major components")
    _h2(doc, "4.1 Data layer")
    _p(
        doc,
        "Single source file dataset.csv with daily rows per cryptocurrency: open, high, low, "
        "close, volume, marketCap, timestamps, crypto_name, and date.",
    )
    _h2(doc, "4.2 Processing and modeling layer")
    _bullets(
        doc,
        [
            "Data preparation: type coercion, missing-value handling, chronological ordering per coin.",
            "Feature engineering: rolling volatility, moving-average ratios, Bollinger-style width, "
            "range and liquidity-style ratios, log market cap, categorical encoding of coin.",
            "Training: scikit-learn HistGradientBoostingRegressor with time-ordered train/test split.",
            "Artifacts: serialized model bundle (joblib), JSON metrics, feature name list.",
        ],
    )
    _h2(doc, "4.3 Presentation layer")
    _p(
        doc,
        "Streamlit application displays holdout metrics, a sample actual-versus-predicted view, "
        "and a latest-row prediction for a user-selected cryptocurrency.",
    )

    _h1(doc, "5. Design constraints")
    _bullets(
        doc,
        [
            "No production cloud deployment in scope; local execution only.",
            "Features must not use future prices; the target is defined using future returns for supervised learning only.",
            "Evaluation uses a chronological holdout to reduce optimistic bias from random shuffling.",
        ],
    )

    _h1(doc, "6. Non-functional considerations")
    _bullets(
        doc,
        [
            "Reproducibility: fixed random seed for the estimator where applicable.",
            "Maintainability: modular layout under src/ with separate data, feature, and training logic.",
        ],
    )

    return doc


def build_lld() -> Document:
    doc = Document()
    _set_normal_style(doc)
    _title(doc, "Low-Level Design (LLD)")
    _p(doc, "Cryptocurrency Volatility Prediction — Component breakdown")
    _p(doc, "Document version: 1.0")

    _h1(doc, "1. Module map")
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Module / file"
    hdr[1].text = "Responsibility"
    hdr[2].text = "Key outputs"
    rows = [
        ("src/paths.py", "Project root paths, artifact and report directories.", "MODEL_PATH, FIGURES, etc."),
        ("src/data_prep.py", "load_raw(), build_modeling_dataframe().", "Clean DataFrame, LabelEncoder, crypto_id."),
        ("src/features.py", "add_features(), feature_columns(), target computation.", "Engineered columns + target_vol."),
        ("src/train.py", "time_ordered_split(), train_and_save().", "joblib model, metrics.json, feature_names.json."),
        ("src/eda.py", "Summary statistics and matplotlib/seaborn figures.", "PNG plots, eda_summary_stats.csv."),
        ("src/predict.py", "load_bundle(), helpers for inference.", "In-memory model dict."),
        ("streamlit_app.py", "Cached data load, UI layout, prediction display.", "Browser UI."),
    ]
    for a, b, c in rows:
        row = table.add_row().cells
        row[0].text = a
        row[1].text = b
        row[2].text = c

    _h1(doc, "2. Data structures")
    _h2(doc, "2.1 Raw dataframe columns")
    _p(
        doc,
        "open, high, low, close, volume, marketCap, timestamp, crypto_name, date (UTC-parsed).",
    )
    _h2(doc, "2.2 Modeled feature vector (before crypto_id)")
    _bullets(
        doc,
        [
            "log_ret — daily log return on close.",
            "roll_vol_5, roll_vol_10, roll_vol_20 — rolling standard deviation of log_ret.",
            "close_ma7_ratio, close_ma20_ratio — close relative to rolling mean minus one.",
            "bb_width — Bollinger band width normalized by middle band.",
            "hl_range_pct — (high − low) / close.",
            "co_range_pct — |close − open| / open.",
            "vol_mcap_ratio — volume / marketCap (guarded for zero cap).",
            "log_mcap — log(1 + max(marketCap, 0)).",
            "crypto_id — integer label from sklearn.preprocessing.LabelEncoder.",
        ],
    )
    _h2(doc, "2.3 Target")
    _p(
        doc,
        "target_vol: population standard deviation of the next five daily log returns "
        "(forward realized volatility), computed per cryptocurrency time series.",
    )

    _h1(doc, "3. Algorithms and parameters")
    _h2(doc, "3.1 Train/test split")
    _p(
        doc,
        "Rows are sorted by date ascending. The first 85% of rows form training; the last "
        "15% form the holdout test set (global chronological split across the stacked panel).",
    )
    _h2(doc, "3.2 Regressor")
    _p(
        doc,
        "HistGradientBoostingRegressor: max_depth=8, learning_rate=0.06, max_iter=200, "
        "min_samples_leaf=40, l2_regularization=1e-3, random_state=42.",
    )

    _h1(doc, "4. File artifacts")
    _bullets(
        doc,
        [
            "artifacts/volatility_model.joblib — dict with model, label_encoder, feature_columns.",
            "artifacts/metrics.json — RMSE, MAE, R², row counts, target description.",
            "artifacts/feature_names.json — ordered feature list for inference.",
        ],
    )

    return doc


def build_pipeline() -> Document:
    doc = Document()
    _set_normal_style(doc)
    _title(doc, "Pipeline Architecture")
    _p(doc, "Data flow: preprocessing to prediction")
    _p(doc, "Document version: 1.0")

    _h1(doc, "1. End-to-end flow")
    _p(
        doc,
        "The following steps run in order for training and batch evaluation. The Streamlit "
        "app reuses the same feature pipeline for consistency with training.",
    )
    steps = [
        "Ingest: Read dataset.csv from the project root (pandas.read_csv, index column dropped).",
        "Clean: Parse date as UTC; coerce numeric columns; drop rows missing OHLC or date; require close > 0; sort by crypto_name and date.",
        "Feature & target: For each crypto_name group, compute log returns, rolling statistics, "
        "Bollinger width, range ratios, volume-to-market-cap ratio, log market cap, and forward "
        "five-day realized volatility target.",
        "Filter: Drop rows with NaN in any modeled feature or in target_vol.",
        "Encode: Fit LabelEncoder on crypto_name; add crypto_id column.",
        "Split: Sort by date; allocate last 15% of rows to test; remainder to train.",
        "Train: Fit HistGradientBoostingRegressor on training matrix X and vector y (target_vol).",
        "Evaluate: Predict on test; compute RMSE, MAE, R²; write metrics.json.",
        "Persist: joblib.dump model bundle; write feature_names.json.",
        "Optional EDA: python -m src.eda writes reports/figures and summary CSV.",
        "App: streamlit_app.py loads bundle, rebuilds modeling dataframe (cached), repeats split logic for display, runs predict for UI.",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}", style="List Number")

    _h1(doc, "2. Logical diagram (textual)")
    _p(
        doc,
        "[ dataset.csv ] → [ data_prep: load_raw ] → [ features: add_features ] → "
        "[ dropna + LabelEncoder ] → [ train: split + fit ] → [ artifacts: .joblib + .json ]",
    )
    _p(
        doc,
        "Parallel branch: same featurized frame → [ eda ] → [ reports/figures ]. "
        "Serving branch: artifacts + featurized frame → [ Streamlit ] → user browser.",
    )

    _h1(doc, "3. Dependencies")
    _bullets(
        doc,
        [
            "Python 3.x",
            "pandas, numpy, scikit-learn, joblib",
            "matplotlib, seaborn (EDA)",
            "streamlit (local UI)",
        ],
    )

    return doc


def build_final_report(metrics: dict) -> Document:
    doc = Document()
    _set_normal_style(doc)
    _title(doc, "Final Report")
    _p(doc, "Cryptocurrency Volatility Prediction")
    _p(doc, "Document version: 1.0")

    _h1(doc, "1. Executive summary")
    _p(
        doc,
        "This project builds a supervised regression model that estimates forward five-day "
        "realized volatility from daily OHLCV and market capitalization history across many "
        "cryptocurrencies. A gradient boosting regressor is trained with a chronological "
        "holdout split; performance is reported with RMSE, MAE, and R². A Streamlit app "
        "exposes metrics and single-point predictions for exploration.",
    )

    _h1(doc, "2. Dataset and preprocessing")
    _p(
        doc,
        "The dataset provides daily observations per coin with open, high, low, close, volume, "
        "and marketCap. Rows with invalid or missing core fields are removed. Each series is "
        "processed in date order so rolling features respect time.",
    )

    _h1(doc, "3. Methodology")
    _h2(doc, "3.1 Target definition")
    _p(
        doc,
        "The target is the standard deviation of the next five daily log returns (forward "
        "realized volatility). This aligns with the goal of anticipating near-term turbulence "
        "while keeping input features based only on information available through the current day.",
    )
    _h2(doc, "3.2 Features")
    _p(
        doc,
        "Inputs combine momentum and volatility history (rolling volatilities, MA distance), "
        "Bollinger bandwidth as a volatility envelope proxy, intraday range and body size "
        "relative to price, a simple liquidity proxy (volume divided by market cap), log "
        "market cap scale, and an identifier for which asset the row belongs to.",
    )
    _h2(doc, "3.3 Model and validation")
    _p(
        doc,
        "HistGradientBoostingRegressor was chosen for strong nonlinear fitting with efficient "
        "handling of numeric features. Evaluation uses the last 15% of rows by date as a holdout "
        "set to mimic deployment on more recent market conditions.",
    )

    _h1(doc, "4. Results")
    if metrics:
        _p(
            doc,
            f"Latest training run metrics on the chronological holdout: RMSE = {metrics.get('rmse', 0):.6f}, "
            f"MAE = {metrics.get('mae', 0):.6f}, R² = {metrics.get('r2', 0):.4f}. "
            f"Training rows: {metrics.get('n_train', 'N/A')}; test rows: {metrics.get('n_test', 'N/A')}.",
        )
        _p(doc, f"Target definition recorded in metrics: {metrics.get('target', '')}.")
    else:
        _p(
            doc,
            "Run python -m src.train to generate artifacts/metrics.json; this section will "
            "then match your local numbers. Placeholder: metrics not found at build time.",
        )

    _h1(doc, "5. Key insights")
    _bullets(
        doc,
        [
            "Volatility is inherently difficult to forecast; moderate R² is expected without external signals or alternative data.",
            "Rolling historical volatility and range-based features provide a sensible baseline for short-horizon realized vol.",
            "Time-ordered evaluation gives a more realistic sense of performance than a random split.",
        ],
    )

    _h1(doc, "6. Deliverables checklist")
    _bullets(
        doc,
        [
            "Source code under src/, streamlit_app.py, scripts, requirements.txt.",
            "Trained model and metrics under artifacts/ (after training).",
            "EDA plots and summary under reports/.",
            "Word documents: HLD, LLD, Pipeline Architecture, EDA Report, Final Report (this folder).",
        ],
    )

    _h1(doc, "7. Future work")
    _bullets(
        doc,
        [
            "Per-asset or grouped time-series cross-validation.",
            "Hyperparameter search (grid or Bayesian) and optional regularization tuning.",
            "Additional indicators or multi-horizon targets; comparison with naive GARCH-style baselines.",
        ],
    )

    return doc


def build_eda_report() -> Document:
    doc = Document()
    _set_normal_style(doc)
    _title(doc, "Exploratory Data Analysis (EDA) Report")
    _p(doc, "Cryptocurrency historical prices — volatility project")
    _p(doc, "Document version: 1.0")

    _h1(doc, "1. Objectives")
    _p(
        doc,
        "Exploratory analysis confirms data quality, reveals distributional properties of "
        "engineered features and the volatility target, and surfaces correlations that inform "
        "model choice. Automated outputs are produced by running: python -m src.eda",
    )

    _h1(doc, "2. Outputs generated")
    _bullets(
        doc,
        [
            "reports/eda_summary_stats.csv — descriptive statistics for numeric features and target_vol.",
            "reports/figures/correlation_heatmap.png — linear association among features and target.",
            "reports/figures/close_trends_sample.png — close price paths for a sample of liquid names.",
            "reports/figures/target_distribution.png — histogram of forward five-day realized volatility.",
        ],
    )

    _h1(doc, "3. How to interpret the visuals")
    _h2(doc, "3.1 Correlation heatmap")
    _p(
        doc,
        "Strong positive correlation between rolling volatility measures and the forward "
        "volatility target is expected. Weaker correlations suggest room for nonlinear models.",
    )
    _h2(doc, "3.2 Price trends")
    _p(
        doc,
        "Sample series illustrate long-horizon growth and drawdowns; non-stationarity motivates "
        "features based on returns and ratios rather than raw price levels alone.",
    )
    _h2(doc, "3.3 Target distribution")
    _p(
        doc,
        "The target is right-skewed with a concentration of moderate volatility and a tail of "
        "high-volatility episodes, consistent with crisis and hype periods in crypto markets.",
    )

    _h1(doc, "4. Summary statistics")
    _p(
        doc,
        "Refer to eda_summary_stats.csv for count, mean, standard deviation, min, quartiles, and "
        "max for each column used in modeling. Regenerate after any change to dataset.csv or "
        "feature definitions.",
    )

    return doc


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    metrics = _load_metrics()

    docs = [
        ("HLD_Cryptocurrency_Volatility_Prediction.docx", build_hld()),
        ("LLD_Cryptocurrency_Volatility_Prediction.docx", build_lld()),
        ("Pipeline_Architecture.docx", build_pipeline()),
        ("EDA_Report.docx", build_eda_report()),
        ("Final_Report.docx", build_final_report(metrics)),
    ]
    for name, d in docs:
        path = DOCS_DIR / name
        d.save(path)
        print("Wrote", path)


if __name__ == "__main__":
    main()
